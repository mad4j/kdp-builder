#!/usr/bin/env python3
"""
KDP Builder - Markdown to DOCX Converter

This script converts Markdown files with style attributes to DOCX format,
using YAML files for style definitions and document layout.
"""

import argparse
import re
import sys
from pathlib import Path
from typing import Dict, List, Any, Tuple

import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class StyleDefinition:
    """Represents a style definition from YAML."""
    
    def __init__(self, style_data: Dict[str, Any]):
        self.font_name = style_data.get('font_name', 'Arial')
        self.font_size = style_data.get('font_size', 11)
        self.bold = style_data.get('bold', False)
        self.italic = style_data.get('italic', False)
        self.underline = style_data.get('underline', False)
        self.color = style_data.get('color')
        self.alignment = style_data.get('alignment', 'left')
        self.space_before = style_data.get('space_before', 0)
        self.space_after = style_data.get('space_after', 0)


class LayoutDefinition:
    """Represents document layout from YAML."""
    
    # Conversion factor: 1 inch = 25.4 mm
    MM_TO_INCHES = 1.0 / 25.4
    
    def __init__(self, layout_data: Dict[str, Any]):
        # Get the unit for dimensions (default to inches for backward compatibility)
        self.unit = layout_data.get('unit', 'inches').lower()
        
        # Validate unit
        if self.unit not in ['inches', 'mm']:
            raise ValueError(f"Invalid unit '{self.unit}'. Must be 'inches' or 'mm'.")
        
        # Get dimension values and convert to inches if necessary
        self.page_width = self._convert_to_inches(layout_data.get('page_width', 8.5))
        self.page_height = self._convert_to_inches(layout_data.get('page_height', 11))
        self.margin_top = self._convert_to_inches(layout_data.get('margin_top', 1.0))
        self.margin_bottom = self._convert_to_inches(layout_data.get('margin_bottom', 1.0))
        self.margin_left = self._convert_to_inches(layout_data.get('margin_left', 1.0))
        self.margin_right = self._convert_to_inches(layout_data.get('margin_right', 1.0))
        self.header_text = layout_data.get('header_text')
        self.header_style = layout_data.get('header_style', 'normal')
        self.footer_text = layout_data.get('footer_text')
        self.footer_style = layout_data.get('footer_style', 'normal')
    
    def _convert_to_inches(self, value: float) -> float:
        """Convert a dimension value to inches based on the current unit."""
        if value < 0:
            raise ValueError(f"Dimension values must be positive, got {value}")
        if self.unit == 'mm':
            return value * self.MM_TO_INCHES
        return value


class MarkdownParser:
    """Parses Markdown with custom style attributes."""
    
    # Pattern to match text with style attributes: {text}[style]
    STYLED_TEXT_PATTERN = re.compile(r'\{([^}]+)\}\[([^\]]+)\]')
    
    # Pattern to match Markdown headers
    HEADER_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$')
    
    # Pattern to match page break markers
    PAGEBREAK_PATTERN = re.compile(r'^<<<pagebreak>>>$', re.IGNORECASE)
    
    # Pattern to match index markers: <<<index:term>>>
    INDEX_PATTERN = re.compile(r'^<<<index:(.+)>>>$', re.IGNORECASE)
    
    # Pattern to match table of contents markers: <<<toc>>>
    TOC_PATTERN = re.compile(r'^<<<toc>>>$', re.IGNORECASE)
    
    @staticmethod
    def is_pagebreak(line: str) -> bool:
        """Check if a line is a page break marker."""
        return bool(MarkdownParser.PAGEBREAK_PATTERN.match(line.strip()))
    
    @staticmethod
    def is_index(line: str) -> Tuple[bool, str]:
        """
        Check if a line is an index marker.
        
        Returns (is_index, term) where term is the index entry text.
        """
        match = MarkdownParser.INDEX_PATTERN.match(line.strip())
        if match:
            return (True, match.group(1))
        return (False, '')
    
    @staticmethod
    def is_toc(line: str) -> bool:
        """Check if a line is a table of contents marker."""
        return bool(MarkdownParser.TOC_PATTERN.match(line.strip()))
    
    @staticmethod
    def parse_line(line: str) -> List[Tuple[str, str]]:
        """
        Parse a line of Markdown into text segments with their styles.
        
        Returns a list of tuples: [(text, style_name), ...]
        If no style is specified, style_name is 'normal'.
        """
        segments = []
        pos = 0
        
        # Check if line is a header
        header_match = MarkdownParser.HEADER_PATTERN.match(line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            style_name = f'heading{level}'
            return [(text, style_name)]
        
        # Process styled text and regular text
        for match in MarkdownParser.STYLED_TEXT_PATTERN.finditer(line):
            # Add any text before this match as normal text
            if match.start() > pos:
                normal_text = line[pos:match.start()]
                if normal_text.strip():
                    segments.append((normal_text, 'normal'))
            
            # Add the styled text
            text = match.group(1)
            style = match.group(2)
            segments.append((text, style))
            pos = match.end()
        
        # Add any remaining text as normal
        if pos < len(line):
            remaining = line[pos:]
            if remaining.strip():
                segments.append((remaining, 'normal'))
        
        # If no segments were found, treat the whole line as normal text
        if not segments and line.strip():
            segments.append((line, 'normal'))
        
        return segments


class DocxBuilder:
    """Builds a DOCX document from parsed Markdown with styles."""
    
    def __init__(self, styles: Dict[str, StyleDefinition], layout: LayoutDefinition):
        self.document = Document()
        self.styles = styles
        self.layout = layout
        # Create a default style if 'normal' is not defined
        if 'normal' not in self.styles:
            self.styles['normal'] = StyleDefinition({})
        self._apply_layout()
    
    @staticmethod
    def _add_field(run, field_name: str):
        """Add a field code to a run."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_name

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    @staticmethod
    def _add_page_number_field(run):
        """Add page number field to a run."""
        DocxBuilder._add_field(run, "PAGE")
    
    @staticmethod
    def _add_page_count_field(run):
        """Add total page count field to a run."""
        DocxBuilder._add_field(run, "NUMPAGES")
    
    def _add_text_with_fields(self, paragraph, text: str, style_def: StyleDefinition):
        """Add text to a paragraph with support for {page} and {total} field codes."""
        if not text:
            return
        
        # Split the text into regular text and placeholder tokens
        # Pattern matches {page} or {total} placeholders
        parts = re.split(r'(\{page\}|\{total\})', text)
        
        # Add runs for each part
        for part in parts:
            if not part:  # Skip empty strings
                continue
            
            run = paragraph.add_run(part if part not in ['{page}', '{total}'] else '')
            self._apply_style_to_run(run, style_def)
            
            if part == '{page}':
                self._add_page_number_field(run)
            elif part == '{total}':
                self._add_page_count_field(run)
    
    def _apply_layout(self):
        """Apply layout settings to the document."""
        sections = self.document.sections
        for section in sections:
            section.page_width = Inches(self.layout.page_width)
            section.page_height = Inches(self.layout.page_height)
            section.top_margin = Inches(self.layout.margin_top)
            section.bottom_margin = Inches(self.layout.margin_bottom)
            section.left_margin = Inches(self.layout.margin_left)
            section.right_margin = Inches(self.layout.margin_right)
    
    def _apply_header_footer(self):
        """Apply header and footer to the document."""
        sections = self.document.sections
        for section in sections:
            # Add header if specified
            if self.layout.header_text:
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                # Clear any existing content
                header_para.clear()
                
                # Apply style to header
                style_def = self.styles.get(self.layout.header_style, self.styles['normal'])
                header_para.alignment = self._get_alignment(style_def.alignment)
                
                # Add text with support for {page} and {total} tags
                self._add_text_with_fields(header_para, self.layout.header_text, style_def)
            
            # Add footer if specified
            if self.layout.footer_text:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                # Clear any existing content
                footer_para.clear()
                
                # Apply style to footer
                style_def = self.styles.get(self.layout.footer_style, self.styles['normal'])
                footer_para.alignment = self._get_alignment(style_def.alignment)
                
                # Add text with support for {page} and {total} tags
                self._add_text_with_fields(footer_para, self.layout.footer_text, style_def)
    
    def _get_alignment(self, alignment_str: str):
        """Convert alignment string to DOCX alignment constant."""
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return alignment_map.get(alignment_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)
    
    def _parse_color(self, color_str: str) -> RGBColor:
        """Parse color string (hex format like '#FF0000') to RGBColor."""
        if color_str.startswith('#'):
            color_str = color_str[1:]
        
        # Validate color string format
        if len(color_str) != 6:
            raise ValueError(f"Invalid color format: {color_str}. Expected 6 hex characters.")
        
        try:
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
        except ValueError:
            raise ValueError(f"Invalid hex color: {color_str}. Must contain only hex digits.")
        
        return RGBColor(r, g, b)
    
    def _apply_style_to_run(self, run, style_def: StyleDefinition):
        """Apply a style definition to a text run."""
        run.font.name = style_def.font_name
        run.font.size = Pt(style_def.font_size)
        run.font.bold = style_def.bold
        run.font.italic = style_def.italic
        run.font.underline = style_def.underline
        
        if style_def.color:
            run.font.color.rgb = self._parse_color(style_def.color)
    
    def add_paragraph(self, segments: List[Tuple[str, str]]):
        """Add a paragraph to the document with styled segments."""
        if not segments:
            self.document.add_paragraph()
            return
        
        # Use the style of the first segment to determine paragraph style
        first_style_name = segments[0][1]
        style_def = self.styles.get(first_style_name, self.styles['normal'])
        
        paragraph = self.document.add_paragraph()
        paragraph.alignment = self._get_alignment(style_def.alignment)
        
        if style_def.space_before > 0:
            paragraph.paragraph_format.space_before = Pt(style_def.space_before)
        if style_def.space_after > 0:
            paragraph.paragraph_format.space_after = Pt(style_def.space_after)
        
        # Add each segment with its style
        for text, style_name in segments:
            style_def = self.styles.get(style_name, self.styles['normal'])
            run = paragraph.add_run(text)
            self._apply_style_to_run(run, style_def)
    
    def add_page_break(self):
        """Add a page break to the document."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def add_index_entry(self, term: str):
        """
        Add an index entry marker to the document.
        
        This creates an XE (Index Entry) field that marks the term for inclusion
        in a document index. The field is hidden in the document but can be used
        to generate an index in Word.
        
        Args:
            term: The term to add to the index
        """
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        
        # Create XE field for index entry
        # Field code format: XE "term"
        # In Word field codes, quotes are escaped by doubling them
        escaped_term = term.replace('"', '""')
        field_code = f'XE "{escaped_term}"'
        
        # Add the field using the existing _add_field method
        DocxBuilder._add_field(run, field_code)
        
        # Make the paragraph hidden so the XE field doesn't show in the document
        # XE fields are typically hidden in Word documents
        pPr = paragraph._element.get_or_add_pPr()
        vanish = OxmlElement('w:vanish')
        pPr.append(vanish)
    
    def add_toc(self):
        """
        Add a table of contents field to the document.
        
        This creates a TOC field that can be updated in Word to generate
        a table of contents based on the document headings.
        """
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        
        # Create TOC field
        # Field code format: TOC \o "1-3" \h \z \u
        # \o "1-3" = include heading levels 1-3
        # \h = make TOC entries hyperlinks
        # \z = hide tab leader and page numbers in Web Layout view
        # \u = use applied paragraph outline level
        field_code = 'TOC \\o "1-3" \\h \\z \\u'
        
        # Add the field using the existing _add_field method
        DocxBuilder._add_field(run, field_code)
    
    def save(self, output_path: str):
        """Save the document to a file."""
        self.document.save(output_path)


def load_styles(yaml_path: str) -> Dict[str, StyleDefinition]:
    """Load style definitions from YAML file."""
    with open(yaml_path, 'r', encoding='utf-8') as f:
        styles_data = yaml.safe_load(f)
    
    styles = {}
    for style_name, style_data in styles_data.get('styles', {}).items():
        styles[style_name] = StyleDefinition(style_data)
    
    return styles


def load_layout(yaml_path: str) -> LayoutDefinition:
    """Load layout definition from YAML file."""
    with open(yaml_path, 'r', encoding='utf-8') as f:
        layout_data = yaml.safe_load(f)
    
    return LayoutDefinition(layout_data.get('layout', {}))


def convert_markdown_to_docx(markdown_path: str, styles_path: str, 
                             layout_path: str, output_path: str):
    """
    Convert a Markdown file to DOCX using style and layout definitions.
    
    Args:
        markdown_path: Path to input Markdown file
        styles_path: Path to YAML file with style definitions
        layout_path: Path to YAML file with layout definition
        output_path: Path to output DOCX file
    """
    # Load configuration
    styles = load_styles(styles_path)
    layout = load_layout(layout_path)
    
    # Create document builder
    builder = DocxBuilder(styles, layout)
    
    # Parse and add Markdown content
    with open(markdown_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if MarkdownParser.is_pagebreak(line):  # Page break marker
                builder.add_page_break()
            elif MarkdownParser.is_toc(line):  # Table of contents marker
                builder.add_toc()
            else:
                # Check for index marker
                is_idx, term = MarkdownParser.is_index(line)
                if is_idx:  # Index marker
                    builder.add_index_entry(term)
                elif line.strip():  # Non-empty line
                    segments = MarkdownParser.parse_line(line)
                    builder.add_paragraph(segments)
                else:  # Empty line
                    builder.add_paragraph([])
    
    # Apply header and footer after content is added
    builder._apply_header_footer()
    
    # Save document
    builder.save(output_path)
    print(f"Document saved to: {output_path}")


def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(
        description='Convert Markdown with style attributes to DOCX format.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Example usage:
  %(prog)s -m input.md -s styles.yaml -l layout.yaml -o output.docx
  
The Markdown file supports custom style attributes:
  {text}[style_name] - Apply custom style to text
  # Heading 1 - Standard Markdown heading
        """
    )
    
    parser.add_argument('-m', '--markdown', required=True,
                       help='Input Markdown file')
    parser.add_argument('-s', '--styles', required=True,
                       help='YAML file with style definitions')
    parser.add_argument('-l', '--layout', required=True,
                       help='YAML file with layout definition')
    parser.add_argument('-o', '--output', required=True,
                       help='Output DOCX file')
    
    args = parser.parse_args()
    
    # Validate input files exist
    for file_path, name in [(args.markdown, 'Markdown'),
                            (args.styles, 'Styles'),
                            (args.layout, 'Layout')]:
        if not Path(file_path).exists():
            print(f"Error: {name} file not found: {file_path}", file=sys.stderr)
            sys.exit(1)
    
    try:
        convert_markdown_to_docx(args.markdown, args.styles, 
                                args.layout, args.output)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
