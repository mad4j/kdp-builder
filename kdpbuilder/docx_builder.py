from __future__ import annotations

import re
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .definitions import LayoutDefinition, StyleDefinition


class DocxBuilder:
    """Builds a DOCX document from parsed Markdown with styles."""

    def __init__(self, styles: Dict[str, StyleDefinition], layout: LayoutDefinition):
        self.document = Document()
        self.styles = styles
        self.layout = layout
        self._bookmark_id_counter = 0

        if "normal" not in self.styles:
            self.styles["normal"] = StyleDefinition()

        self._enable_update_fields_on_open()
        self._apply_layout()

    def _enable_update_fields_on_open(self) -> None:
        """Ask Word to update fields (TOC, PAGE, etc.) when opening the file."""
        settings = self.document.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def _get_next_bookmark_id(self) -> str:
        self._bookmark_id_counter += 1
        return str(self._bookmark_id_counter)

    @staticmethod
    def _add_field(run, field_name: str):
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        fld_char_begin.set(qn("w:dirty"), "true")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = field_name

        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_sep)
        # Result will be populated by Word when fields are updated
        run._r.append(fld_char_end)

    @staticmethod
    def _add_page_number_field(run):
        DocxBuilder._add_field(run, "PAGE")

    @staticmethod
    def _add_page_count_field(run):
        DocxBuilder._add_field(run, "NUMPAGES")

    def _add_text_with_fields(self, paragraph, text: str, style_def: StyleDefinition):
        if not text:
            return

        parts = re.split(r"(\{page\}|\{total\})", text)
        for part in parts:
            if not part:
                continue

            run = paragraph.add_run(part if part not in ["{page}", "{total}"] else "")
            self._apply_style_to_run(run, style_def)

            if part == "{page}":
                self._add_page_number_field(run)
            elif part == "{total}":
                self._add_page_count_field(run)

    def _apply_layout(self):
        for section in self.document.sections:
            section.page_width = Inches(self.layout.page_width)
            section.page_height = Inches(self.layout.page_height)
            section.top_margin = Inches(self.layout.margin_top)
            section.bottom_margin = Inches(self.layout.margin_bottom)
            section.left_margin = Inches(self.layout.margin_left)
            section.right_margin = Inches(self.layout.margin_right)

    def apply_header_footer(self):
        """Apply header and footer to the document (after content is added)."""
        # Do not show header/footer on the very first page of the document.
        # In Word this is controlled per-section; we apply it only to the first section
        # so subsequent sections (if any) keep their normal behavior.
        if self.document.sections and (self.layout.header_text or self.layout.footer_text):
            first_section = self.document.sections[0]
            first_section.different_first_page_header_footer = True

            if self.layout.header_text:
                first_page_header = first_section.first_page_header
                first_page_header.is_linked_to_previous = False
                header_para = (
                    first_page_header.paragraphs[0]
                    if first_page_header.paragraphs
                    else first_page_header.add_paragraph()
                )
                header_para.clear()

            if self.layout.footer_text:
                first_page_footer = first_section.first_page_footer
                first_page_footer.is_linked_to_previous = False
                footer_para = (
                    first_page_footer.paragraphs[0]
                    if first_page_footer.paragraphs
                    else first_page_footer.add_paragraph()
                )
                footer_para.clear()

        for section in self.document.sections:
            if self.layout.header_text:
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.clear()

                style_def = self.styles.get(self.layout.header_style, self.styles["normal"])
                header_para.alignment = self._get_alignment(style_def.alignment)
                self._add_text_with_fields(header_para, self.layout.header_text, style_def)

            if self.layout.footer_text:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.clear()

                style_def = self.styles.get(self.layout.footer_style, self.styles["normal"])
                footer_para.alignment = self._get_alignment(style_def.alignment)
                self._add_text_with_fields(footer_para, self.layout.footer_text, style_def)

    def _get_alignment(self, alignment_str: str):
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return alignment_map.get(str(alignment_str).lower(), WD_ALIGN_PARAGRAPH.LEFT)

    def _parse_color(self, color_str: str) -> RGBColor:
        if color_str.startswith("#"):
            color_str = color_str[1:]

        if len(color_str) != 6:
            raise ValueError(
                f"Invalid color format: {color_str}. Expected 6 hex characters."
            )

        try:
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
        except ValueError as exc:
            raise ValueError(
                f"Invalid hex color: {color_str}. Must contain only hex digits."
            ) from exc

        return RGBColor(r, g, b)

    def _apply_style_to_run(self, run, style_def: StyleDefinition):
        run.font.name = style_def.font_name
        run.font.size = Pt(style_def.font_size)
        run.font.bold = style_def.bold
        run.font.italic = style_def.italic
        run.font.underline = style_def.underline

        if style_def.color:
            run.font.color.rgb = self._parse_color(style_def.color)

    def add_paragraph(self, segments: List[Tuple[str, str, str]], auto_bookmark: str | None = None):
        if not segments:
            self.document.add_paragraph()
            return

        first_style_name = segments[0][1]
        style_def = self.styles.get(first_style_name, self.styles["normal"])

        paragraph = self.document.add_paragraph()

        heading_match = re.match(r"^heading(\d+)$", str(first_style_name).strip(), re.IGNORECASE)
        if heading_match:
            level = int(heading_match.group(1))
            try:
                paragraph.style = f"Heading {level}"
            except KeyError:
                # If the built-in heading style isn't available, keep default style.
                pass
        paragraph.alignment = self._get_alignment(style_def.alignment)

        if style_def.space_before > 0:
            paragraph.paragraph_format.space_before = Pt(style_def.space_before)
        if style_def.space_after > 0:
            paragraph.paragraph_format.space_after = Pt(style_def.space_after)

        bookmark_id = None
        if auto_bookmark:
            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_id = self._get_next_bookmark_id()
            bookmark_start.set(qn("w:id"), bookmark_id)
            bookmark_start.set(qn("w:name"), auto_bookmark)
            paragraph._p.insert(0, bookmark_start)

        for text, style_name, link_target in segments:
            style_def = self.styles.get(style_name, self.styles["normal"])
            if link_target:
                self._add_hyperlink(paragraph, text, link_target, style_def)
            else:
                run = paragraph.add_run(text)
                self._apply_style_to_run(run, style_def)

        if auto_bookmark and bookmark_id is not None:
            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), bookmark_id)
            paragraph._p.append(bookmark_end)

    def add_page_break(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def add_index_entry(self, term: str):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()

        escaped_term = term.replace('"', '""')
        field_code = f'XE "{escaped_term}"'
        DocxBuilder._add_field(run, field_code)

        pPr = paragraph._element.get_or_add_pPr()
        vanish = OxmlElement("w:vanish")
        pPr.append(vanish)

    def add_toc(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        field_code = 'TOC \\o "1-3" \\h \\z \\u'
        DocxBuilder._add_field(run, field_code)

    def add_bookmark(self, name: str):
        paragraph = self.document.add_paragraph()

        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_id = self._get_next_bookmark_id()
        bookmark_start.set(qn("w:id"), bookmark_id)
        bookmark_start.set(qn("w:name"), name)

        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), bookmark_id)

        paragraph._p.append(bookmark_start)
        paragraph._p.append(bookmark_end)

    @staticmethod
    def _add_hyperlink(paragraph, text: str, bookmark: str, style_def: StyleDefinition):
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)

        if style_def.font_name:
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), style_def.font_name)
            rFonts.set(qn("w:hAnsi"), style_def.font_name)
            rPr.append(rFonts)

        if style_def.font_size:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(style_def.font_size * 2))
            rPr.append(sz)

        if style_def.bold:
            rPr.append(OxmlElement("w:b"))

        if style_def.italic:
            rPr.append(OxmlElement("w:i"))

        new_run.append(rPr)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def save(self, output_path: str):
        self.document.save(output_path)
